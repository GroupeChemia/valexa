from kaleido.scopes.plotly import PlotlyScope
import plotly.graph_objects as go
from zipfile import ZipFile
import os
from docx import Document
from docx.shared import Mm
import shutil
import win32com.client as client
import time
from valexa.helper import roundsf
from datetime import date




###   MAIN   ###
def generateWord(**listProfiles):
    doc = Document("filesReport/TemplateReport.docx")
    version = str(date.today().year)[:-2] + chr(date.today().month + 64) + str(date.today().day)

    commands = {
        "DATEGENERATIONdoc": str(date.today()),
        "AUTEURdoc": "[AUTEUR]",
        "PCNAMEdoc": "PC-MAT-[###]",
        "NUMEROdoc": "PC-MAT-[###]-VAL",
        "VERSIONdoc": version,
        "EXPIRATIONdoc": "[EXPIRATION]",

        "commandPROFILE": listProfiles
    }
    for old, new in commands.items():
        replaceInDocx(doc, old, new)

    doc.save('filesReport/RapportValidation.docx')


def replaceInDocx(docObj, old, new):
    for p in docObj.paragraphs:
        if p.text.find(old) > -1:
            inline = p.runs
            for i in range(len(inline)):
                if inline[i].text.find(old) > -1:
                    if old == "commandPROFILE":
                        inline[i].text = inline[i].text.replace(old, "")
                        generateProfiles(docObj, **new)
                    else:
                        inline[i].text = inline[i].text.replace(old, new)
    for table in docObj.tables:
        for row in table.rows:
            for cell in row.cells:
                replaceInDocx(cell, old, new)
    if docObj.__class__.__name__ == "Document":
        for s in docObj.sections:
            replaceInDocx(s.header, old, new)


def fillTable(table, headers, valuesCol):
    for c in range(len(headers)):
        table.cell(0, c).text = headers[c]
        table.cell(0, c).paragraphs[0].runs[0].font.bold = True
    for l in range(1, len(table.rows)):  #début ligne 1
        for c in range(len(table.columns)):
            table.cell(l, c).text = valuesCol[c][l-1]


def createGraphs(id, data, layout):  # data=list, layout=dict
    """
    Create .png with graphs from the Profile vue

    :param str id:      str id of the specific profile
    :param list data:   list of dict of properties passed to the constructor of the specified trace type
    :param dict layout: dict of properties passed to the Layout constructor
    """
    layout['title']['x'] = 0.5
    layout['title']['y'] = 0.92
    layout['legend']['x'] = 0.9
    layout['legend']['y'] = 1
    scope = PlotlyScope()
    fig = go.Figure(data=data, layout=layout)
    with open("filesReport/profiles/" + id + ".png", "wb") as f:
        f.write(scope.transform(fig, format="png"))


def generateProfiles(docObj, **listProfiles):
    docObj.add_heading('Profiles', 1)
    for profile in listProfiles['profiles']:
        generateProfile(docObj, **profile)
        docObj.add_page_break()
    ## TODO: ajout page references/conclusion


def generateProfile(docObj, **data):
    idProfile = data['compound_name'] + str(data['id'])
    modelInfo = data['model_info']

    createGraphs(id="figPROFILE_" + idProfile,
                 data=data['graphs']['profile']['data'],
                 layout=data['graphs']['profile']['layout'])
    createGraphs(id="figLINEARITY_" + idProfile,
                 data=data['graphs']['linearity']['data'],
                 layout=data['graphs']['linearity']['layout'])
    if 'correction' in data['graphs']:
        createGraphs(id="figCORRECTION_" + idProfile,
                     data=data['graphs']['correction']['data'],
                     layout=data['graphs']['correction']['layout'])

    docObj.add_heading('Composé ' + data['compound_name'], 2)

    ###   SUMMARY   ###
    docObj.add_heading('Résumé', 3)
    table = docObj.add_table(8, 2)
    table.style = 'Table Grid'
    fillTable(table, ["Parameter", "Value"], [
        ["LOD", "LOQ Min", "LOQ Max", "Correction", "Recouvrement Moyen", "Tolerance", "Acceptance"],
        [
            f"""{modelInfo['lod']} {modelInfo['units']} ({
                "" if modelInfo['lod_type'] is None else modelInfo['lod_type']
            })""",
            f"{modelInfo['min_loq']} {modelInfo['units']}",
            f"{modelInfo['max_loq']} {modelInfo['units']}",
            f"""{
                str(modelInfo['correction_factor']) + ("", " (Forced)")[modelInfo['forced_correction_value'] is not None] 
                if modelInfo['has_correction'] else "---"
            }""",
            f"{modelInfo['average_recovery']}",
            f"{modelInfo['tolerance']} %",
            f"""{modelInfo['acceptance']} {modelInfo['units'] + " (Relative)" if modelInfo['absolute_acceptance'] else "% (Absolute)"}""",
        ]
    ])
    docObj.add_picture("filesReport/profiles/figPROFILE_" + idProfile + ".png", width=Mm(150))

    ###   TRUENESS   ###
    docObj.add_heading('Justesse', 3)
    table = docObj.add_table(len(data['levels_info'])+1, 8)
    table.style = 'Table Grid'
    fillTable(table,
              ["Niveau", "Conc.", "Conc. Calc.", "Biais Abs. (%)", "Biais Rel. (%)", "Récupération (%)", "Abs. Tol.", "Rel. Tol."],
              [
                  [str(item) for item in list(range(len(data['levels_info'])))],
                  [str(item['introduced_concentration']) for item in data['levels_info']],
                  [str(item['calculated_concentration']) for item in data['levels_info']],
                  [str(item['bias_abs']) for item in data['bias_info']],
                  [str(item['bias_rel']) for item in data['bias_info']],
                  [str(item['recovery']) for item in data['bias_info']],
                  [f"{item['tolerance_abs_high']}, {item['tolerance_abs_low']}" for item in data['tolerance_info']],
                  [f"{item['tolerance_rel_high']}, {item['tolerance_rel_low']}" for item in data['tolerance_info']],
              ])

    ###  PRECISION REPEATABILITY TABLE  ###
    docObj.add_heading('Fidélité et Répétabilité', 3)
    table = docObj.add_table(len(data['levels_info'])+1, 7)
    table.style = 'Table Grid'
    fillTable(table,
              ["Niveau", "Conc.", "Préc. Inter. Abs.",
               "Préc. Inter. Rel.", "Rep. Abs.", "Rep. Rel.", "Ratio Var."],
              [
                  [str(item) for item in list(range(len(data['levels_info'])))],  # nb ligne tabular
                  [str(item['introduced_concentration']) for item in data['levels_info']],
                  [str(item['intermediate_precision_std']) for item in data['intermediate_precision']],
                  [str(item['intermediate_precision_cv']) for item in data['intermediate_precision']],
                  [str(item['repeatability_std']) for item in data['repeatability_info']],
                  [str(item['repeatability_cv']) for item in data['repeatability_info']],
                  [str(item['ratio_var']) for item in data['misc_stats']]
              ])
    ###  UNCERTAINTY  ###
    docObj.add_heading('Incertitude', 3)
    table = docObj.add_table(len(data['levels_info'])+1, 5)
    table.style = 'Table Grid'
    fillTable(table,
              ["Level", "Conc.", "Conc. Calc.", "Incert. Élargie Abs.", "Incert. Élargie. Rel. (%)"],
              [
                  [str(item) for item in list(range(len(data['levels_info'])))],  # nb ligne tabular
                  [str(item['introduced_concentration']) for item in data['levels_info']],
                  [str(item['calculated_concentration']) for item in data['levels_info']],
                  [str(item['uncertainty_abs']) for item in data['uncertainty_info']],
                  [str(item['uncertainty_pc']) for item in data['uncertainty_info']]
              ])

    ###  LINEARITY  ###
    docObj.add_heading('Linearité', 3)
    table = docObj.add_table(4, 2)
    table.style = 'Table Grid'
    fillTable(table,
              ["Parameter", "Valeur"],
              [
                  ["Pente", "Ordonnée Origine", "R^2"],
                  [
                      str(data['linearity_info']['slope']['value']),
                      str(data['linearity_info']['intercept']['value']),
                      str(data['linearity_info']['rsquared']['value'])
                  ]
              ])
    docObj.add_picture("filesReport/profiles/figLINEARITY_" + idProfile + ".png", width=Mm(150))

    if 'correction' in data['graphs']:
        ###  LINEARITY W/O CORRECTION  ###
        docObj.add_heading('Linearité sans Correction', 3)
        table = docObj.add_table(4, 2)
        table.style = 'Table Grid'
        fillTable(table,
                  ["Parameter", "Valeur"],
                  [
                      ["Pente", "Ordonnée Origine", "R^2"],
                      [
                          str(data['correction_info']['slope']['value']),
                          str(data['correction_info']['intercept']['value']),
                          str(data['correction_info']['rsquared']['value'])
                      ]
                  ])
        docObj.add_picture("filesReport/profiles/figCORRECTION_" + idProfile + ".png", width=Mm(150))

    ###  VALIDATION TABLE  ###
    docObj.add_heading('Données de validation', 3)
    table = docObj.add_table(len(data['validation_data'])+1, 7)
    table.style = 'Table Grid'
    fillTable(table,
              ["Serie", "Niveau", "Conc.", "Réponse", "Conc. Calc.", "Biais Abs.", "Biais Rel."],
              [
                  [str(item['Series']) for item in data['validation_data']],
                  [str(item['Level']) for item in data['validation_data']],
                  [str(roundsf(item['x'],4)) for item in data['validation_data']],
                  [str(roundsf(item['y'],4)) for item in data['validation_data']],
                  [str(roundsf(item['x_calc'],4)) for item in data['validation_data']],
                  [str(item['bias_abs']) for item in data['validation_data']],
                  [str(item['bias_rel']) for item in data['validation_data']],
              ])

    ###  CALIBRATION REGRESSION  ###
    if 'calibration_data' in data:
        generatePartCalibration(docObj, **data)


def generatePartCalibration(docObj, **data):
    idProfile = data['compound_name'] + str(data['id'])

    createGraphs(id="figREGRESSION_" + idProfile,
                 data=data['graphs']['regression']['data'],
                 layout=data['graphs']['regression']['layout'])
    createGraphs(id="figRESIDUALS_" + idProfile,
                 data=data['graphs']['residuals']['data'],
                 layout=data['graphs']['residuals']['layout'])
    createGraphs(id="figRESIDUALSstd_" + idProfile,
                 data=data['graphs']['residuals_std']['data'],
                 layout=data['graphs']['residuals_std']['layout'])


    docObj.add_heading('Régression de calibration', 3)
    table = docObj.add_table(len(data['regression_info'])+1, 3)
    table.style = 'Table Grid'
    fillTable(table,
              ["Série", "Equation", "R^2"],
              [
                  [str(item) for item in list(range(len(data['regression_info'])))],  # nb ligne tabular,
                  [str(item['function_string']) for item in data['regression_info']],
                  [str(item['rsquared']) for item in data['regression_info']]
              ])
    docObj.add_picture("filesReport/profiles/figREGRESSION_" + idProfile + ".png", width=Mm(150))
    docObj.add_picture("filesReport/profiles/figRESIDUALS_" + idProfile + ".png", width=Mm(150))
    docObj.add_picture("filesReport/profiles/figRESIDUALSstd_" + idProfile + ".png", width=Mm(150))

    ###  CALIBRATION TABLE  ###
    docObj.add_heading('Données de calibration', 3)
    table = docObj.add_table(len(data['calibration_data'])+1, 4)
    table.style = 'Table Grid'
    fillTable(table,
              ["Serie", "Niveau", "Concentration", "Réponse"],
              [
                  [str(item['Series']) for item in data['calibration_data']],
                  [str(item['Level']) for item in data['calibration_data']],
                  [str(roundsf(item['x'],4)) for item in data['calibration_data']],
                  [str(roundsf(item['y'],4)) for item in data['calibration_data']],
              ])

def downloadWord():
    """
    Download a .word file of the report in the "Downloads" directory of the PC user
    """
    dirDownloads = os.path.join(os.path.expanduser("~"), "Downloads")
    shutil.copy("filesReport/RapportValidation.docx", dirDownloads + "/RapportValidation" + str(time.time()) + ".docx")

def downloadPdf(isWord):
    """
    Download a .pdf file of the report in the "Downloads" directory of the PC user
    """
    dirDownloads = os.path.join(os.path.expanduser("~"), "Downloads")
    if not isWord:
        downloadWord()  # fichier temporaire accessible a l'appli word
    filepath = dirDownloads + "/RapportValidation.docx"

    word = client.DispatchEx("Word.Application")
    target_path = filepath.replace(".docx", r".pdf")
    word_doc = word.Documents.Open(filepath)
    word_doc.SaveAs(target_path, FileFormat=17)
    word_doc.Close()
    word.Quit()

    if not isWord:
        os.remove(filepath)


def downloadZipGraph():
    """
    Download a .zip file of the images of the report in the "Downloads" directory of the PC user
    """
    file_paths = []
    for root, directories, files in os.walk('filesReport/profiles'):
        for filename in files:
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)
    dirDownloads = os.path.join(os.path.expanduser("~"), "Downloads")
    file_paths.append(dirDownloads + "/RapportValidation.pdf")
    with ZipFile(dirDownloads + '/RapportValidation.zip', 'w') as zip:
        for file in file_paths:
            zip.write(file)